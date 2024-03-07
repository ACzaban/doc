import json
import random
from docx import Document

def lista_taski(json_file, start_index, end_index, step):
    with open(json_file, 'r') as file:
        tasks = json.load(file)
    return tasks[start_index:end_index:step]

def replace_paragraphs_with_random_text(doc_path, taski, date, date2, total_hours):
    doc = Document(doc_path)
    placeholders_task = ['text0', 'text1', 'text2', 'text3', 'text4', 'text5', 'text6', 'text7', 'text8','text9']
    placeholders_hour = ['0h', '1h', '2h', '3h', '4h', '5h', '6h', '7h', '8h', '9h']
    placeholder_date = 'start/m_short/y_short - end/m_short/y_short'
    placeholder_date2 = 'month year'
        
    random_tasks = taski

    random_hours = [random.randint(3, 21) for _ in range(len(random_tasks))]
    while sum(random_hours) != total_hours:
        random_hours = [random.randint(3, 21) for _ in range(len(random_tasks))]

    if len(random_tasks) < 10:
        total_hours_existing = sum(random_hours)
        remaining_tasks = 10 - len(random_tasks)
        remaining_hours = total_hours - total_hours_existing
        
        if remaining_tasks > 0 and remaining_hours > 0:
            factor = remaining_hours / remaining_tasks
            
            for i in range(len(random_hours)):
                if i >= len(random_tasks):
                    random_hours[i] = min(random_hours[i] + factor, 21) 
                    
        for placeholder_task in placeholders_task[len(random_tasks):]:
            for paragraph in doc.paragraphs:
                if placeholder_task in paragraph.text:
                    paragraph.clear()  
                    break


    for placeholder_task, task, hours in zip(placeholders_task, random_tasks, random_hours):
        
        for paragraph in doc.paragraphs:
            if placeholder_task in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder_task, task)
                break  

    for placeholder_hour, hours in zip(placeholders_hour, random_hours):
        for paragraph in doc.paragraphs:
            if placeholder_hour in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder_hour, str(hours) + ' h')
                break  
    
    for paragraph in doc.paragraphs:
        if 'name' in paragraph.text:
            paragraph.text = paragraph.text.replace('name', 'Norbert Sulżycki')    
    
    for paragraph in doc.paragraphs:
         if placeholder_date in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder_date, str(date))
                
    for paragraph in doc.paragraphs:
         if placeholder_date2 in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder_date2, str(date2))
             

    return doc


if __name__ == "__main__":
    document_path = "szablon.docx"
    json_file = "moje.json"
    
    tasks_5_to_50 = lista_taski(json_file, 4, 50, 7)
    tasks_3_50 = lista_taski(json_file, 3, 50, 7)
    tasks_combined_5_50 = tasks_5_to_50 + tasks_5_to_50
    tasks_5_3 = tasks_combined_5_50[:9]
    modified_doc_step_5 = replace_paragraphs_with_random_text(document_path, tasks_5_3, '01/Aug/22 - 31/Aug/22', 'August 2022',80)
    modified_doc_step_5.save("N. Sulżycki 08.2022.docx")
    print("Second document generated successfully.")

    tasks_6_to_50 = lista_taski(json_file, 5, 50, 7)
    tasks_2_50 = lista_taski(json_file, 2, 50, 7)
    tasks_combined_6_50 = tasks_6_to_50 + tasks_2_50
    tasks_6_2 = tasks_combined_6_50[:8]
    modified_doc_step_6 = replace_paragraphs_with_random_text(document_path, tasks_6_2, '01/Sept/22 - 30/Sept/22', 'September 2022',80)
    modified_doc_step_6.save("N. Sulżycki 09.2022.docx")
    print("Second document generated successfully.")

    tasks_7_to_50 = lista_taski(json_file, 6, 50, 7)
    tasks_1_50 = lista_taski(json_file, 1, 50, 7)
    tasks_combined_7_50 = tasks_7_to_50 + tasks_1_50
    tasks_7_1 = tasks_combined_7_50[:10]
    modified_doc_step_7 = replace_paragraphs_with_random_text(document_path, tasks_7_1, '01/Oct/22 - 31/Oct/22', 'October 2022',110)
    modified_doc_step_7.save("N. Sulżycki 10.2022.docx")
    print("Second document generated successfully.")

    tasks_8_to_50 = lista_taski(json_file, 7, 50, 7)
    tasks_1_50 = lista_taski(json_file, 1, 50, 7)
    tasks_combined_8_50 = tasks_8_to_50 + tasks_1_50
    tasks_8_1 = tasks_combined_8_50[:8]
    modified_doc_step_8 = replace_paragraphs_with_random_text(document_path, tasks_8_1, '01/Nov/22 - 30/Nov/22', 'November 2022',80)
    modified_doc_step_8.save("N. Sulżycki 11.2022.docx")
    print("Second document generated successfully.")

    tasks_13_50 = lista_taski(json_file, 12, 50, 7)
    tasks_11_50 = lista_taski(json_file, 11, 50, 7)
    tasks_combined_13_50 = tasks_13_50 + tasks_11_50
    tasks_13_11 = tasks_combined_13_50[:9]
    modified_doc_step_5 = replace_paragraphs_with_random_text(document_path, tasks_13_11, '01/Dec/22 - 31/Dec/22', 'December 2022',93)
    modified_doc_step_5.save("N. Sulżycki 12.2022.docx")
    print("Second document generated successfully.")

    tasks_14_50 = lista_taski(json_file, 13, 50, 7)
    tasks_10_50 = lista_taski(json_file, 10, 50, 7)
    tasks_combined_14_50 = tasks_14_50 + tasks_10_50
    tasks_14_10 = tasks_combined_14_50[:7]
    modified_doc_step_5 = replace_paragraphs_with_random_text(document_path, tasks_14_10, '01/Jan/23 - 31/Jan/23', 'January 2023',80)
    modified_doc_step_5.save("N. Sulżycki 01.2023.docx")
    print("Second document generated successfully.")

    tasks_28_0 = lista_taski(json_file, 27, 0, -7)
    tasks_22_0 = lista_taski(json_file, 22, 0, -7)
    tasks_combined_28_0 = tasks_28_0 + tasks_22_0
    tasks_28_22 = tasks_combined_28_0[:8]
    modified_doc_step = replace_paragraphs_with_random_text(document_path, tasks_28_22, '01/Feb/23 - 28/February/23', 'February 2023',80)
    modified_doc_step.save("N. Sulżycki 02.2023.docx")
    print("2 document generated successfully.")

    tasks_29_0 = lista_taski(json_file, 28, 0, -7)
    tasks_23_0 = lista_taski(json_file, 23, 0, -7)
    tasks_combined_29_0 = tasks_29_0 + tasks_23_0
    tasks_29_23 = tasks_combined_29_0[:7]
    modified_doc_step = replace_paragraphs_with_random_text(document_path, tasks_29_23, '01/Mar/23 - 31/Mar/23', 'March 2023',80)
    modified_doc_step.save("N. Sulżycki 03.2023.docx")
    print("2 document generated successfully.")

    tasks_32_40 = lista_taski(json_file, 32, 40, 1)
    modified_32_40 = replace_paragraphs_with_random_text(document_path, tasks_32_40,'01/Apr/23 - 30/Apr/23', 'April 2023',80)
    modified_32_40.save("N. Sulżycki 04.2023.docx")
    print("Second document generated successfully.")

    tasks_1_8 = lista_taski(json_file, 0, 8, 1)
    modified_1_8 = replace_paragraphs_with_random_text(document_path, tasks_1_8, '01/May/23 - 31/May/23', 'May 2023',89)
    modified_1_8.save("N. Sulżycki 05.2023.docx")
    print("First document generated successfully.")

    tasks_8_16 = lista_taski(json_file, 8, 16, 1)
    modified_8_16 = replace_paragraphs_with_random_text(document_path, tasks_8_16,'01/June/23 - 30/June/23', 'June 2023',123)
    modified_8_16.save("N. Sulżycki 06.2023.docx")
    print("Second document generated successfully.")

    tasks_16_24 = lista_taski(json_file, 16, 24, 1)
    modified_16_24 = replace_paragraphs_with_random_text(document_path, tasks_16_24,'01/July/23 - 31/July/23', 'July 2023',121)
    modified_16_24.save("N. Sulżycki 07.2023.docx")
    print("Second document generated successfully.")

    tasks_24_32 = lista_taski(json_file, 24, 32, 1)
    modified_24_32 = replace_paragraphs_with_random_text(document_path, tasks_24_32,'01/Aug/23 - 31/Aug/23', 'August 2023',112)
    modified_24_32.save("N. Sulżycki 08.2023.docx")
    print("Second document generated successfully.")

    tasks_43_0 = lista_taski(json_file, 42, 0, -7)
    tasks_42_0 = lista_taski(json_file, 41, 0, -7)
    tasks_combined_43_0 = tasks_43_0 + tasks_42_0
    tasks_43_42 = tasks_combined_43_0[:9]
    modified_doc_step = replace_paragraphs_with_random_text(document_path, tasks_43_42, '01/Sept/23 - 30/Sept/23', 'September 2023',99)
    modified_doc_step.save("N. Sulżycki 09.2023.docx")
    print("2 document generated successfully.")

    tasks_42_0 = lista_taski(json_file, 41, 0, -7)
    tasks_38_0 = lista_taski(json_file, 38, 0, -7)
    tasks_combined_42_0 = tasks_42_0 + tasks_38_0
    tasks_42_38 = tasks_combined_42_0[:9]
    modified_doc_step = replace_paragraphs_with_random_text(document_path, tasks_42_38, '01/Oct/23 - 31/Oct/23', 'October 2023',98)
    modified_doc_step.save("N. Sulżycki 10.2023.docx")
    print("2 document generated successfully.")

    tasks_41_0 = lista_taski(json_file, 40, 0, -7)
    tasks_37_0 = lista_taski(json_file, 37, 0, -7)
    tasks_combined_41_0 = tasks_41_0 + tasks_37_0
    tasks_41_37 = tasks_combined_41_0[:10]
    modified_doc_step = replace_paragraphs_with_random_text(document_path, tasks_41_37, '01/Nov/23 - 30/Nov/23', 'November 2023',92)
    modified_doc_step.save("N. Sulżycki 11.2023.docx")
    print("2 document generated successfully.")

    tasks_40_0 = lista_taski(json_file, 39, 0, -7)
    tasks_38_0 = lista_taski(json_file, 38, 0, -7)
    tasks_combined_40_0 = tasks_40_0 + tasks_38_0
    tasks_40_38 = tasks_combined_41_0[:8]
    modified_doc_step = replace_paragraphs_with_random_text(document_path, tasks_40_38, '01/Dec/23 - 31/Dec/23', 'December 2023',12)
    modified_doc_step.save("N. Sulżycki 12.2023.docx")
    print("2 document generated successfully.")






    